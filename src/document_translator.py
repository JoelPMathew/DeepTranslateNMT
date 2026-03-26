"""
Document Translation Module
Handles translation of various document formats (PDF, DOCX, TXT) 
while preserving formatting and structure.
"""
import os
import json
from typing import Dict, List, Tuple, Optional
from pathlib import Path
from dataclasses import dataclass, asdict
from enum import Enum
import re


class DocumentFormat(Enum):
    """Supported document formats"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MARKDOWN = "md"
    JSON = "json"


@dataclass
class TranslatedSegment:
    """Represents a translated text segment"""
    original: str
    translated: str
    start_offset: int
    end_offset: int
    segment_type: str = "text"  # text, heading, footer, etc.
    metadata: Dict = None

    def to_dict(self):
        return asdict(self)


class DocumentParser:
    """Base parser for document formats"""

    @staticmethod
    def detect_format(file_path: str) -> DocumentFormat:
        """Detect document format from file extension"""
        ext = Path(file_path).suffix.lower()
        ext_map = {
            '.pdf': DocumentFormat.PDF,
            '.docx': DocumentFormat.DOCX,
            '.txt': DocumentFormat.TXT,
            '.md': DocumentFormat.MARKDOWN,
            '.json': DocumentFormat.JSON,
        }
        return ext_map.get(ext, DocumentFormat.TXT)

    @staticmethod
    def extract_segments(file_path: str) -> List[TranslatedSegment]:
        """Extract translatable segments from document"""
        fmt = DocumentParser.detect_format(file_path)
        
        if fmt == DocumentFormat.TXT:
            return PlainTextParser.parse(file_path)
        elif fmt == DocumentFormat.MARKDOWN:
            return MarkdownParser.parse(file_path)
        elif fmt == DocumentFormat.JSON:
            return JSONParser.parse(file_path)
        elif fmt == DocumentFormat.PDF:
            return PDFParser.parse(file_path)
        elif fmt == DocumentFormat.DOCX:
            return DocxParser.parse(file_path)
        else:
            raise ValueError(f"Unsupported format: {fmt}")


class PlainTextParser:
    """Parser for plain text files"""

    @staticmethod
    def parse(file_path: str) -> List[TranslatedSegment]:
        """Parse plain text file"""
        segments = []
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split by paragraphs (double newlines) or sentences
        paragraphs = content.split('\n\n')
        offset = 0
        
        for para in paragraphs:
            if para.strip():
                segment = TranslatedSegment(
                    original=para.strip(),
                    translated="",
                    start_offset=offset,
                    end_offset=offset + len(para),
                    segment_type="paragraph"
                )
                segments.append(segment)
            offset += len(para) + 2  # +2 for double newline
        
        return segments

    @staticmethod
    def reconstruct(segments: List[TranslatedSegment]) -> str:
        """Reconstruct document from translated segments"""
        return "\n\n".join(seg.translated for seg in segments if seg.translated)


class MarkdownParser:
    """Parser for Markdown files"""

    @staticmethod
    def parse(file_path: str) -> List[TranslatedSegment]:
        """Parse Markdown file preserving structure"""
        segments = []
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Regex patterns for Markdown elements
        patterns = [
            (r'^(#{1,6})\s+(.+)$', 'heading'),
            (r'^- (.+)$', 'list_item'),
            (r'^> (.+)$', 'blockquote'),
            (r'^(.+)$', 'paragraph'),
        ]
        
        offset = 0
        for line in content.split('\n'):
            for pattern, seg_type in patterns:
                match = re.match(pattern, line)
                if match:
                    text = match.group(2) if seg_type == 'heading' else match.group(1)
                    if text and text.strip():
                        segment = TranslatedSegment(
                            original=text,
                            translated="",
                            start_offset=offset,
                            end_offset=offset + len(line),
                            segment_type=seg_type,
                            metadata={'prefix': match.group(1) if seg_type == 'heading' else ''}
                        )
                        segments.append(segment)
                    break
            offset += len(line) + 1

        return segments

    @staticmethod
    def reconstruct(segments: List[TranslatedSegment]) -> str:
        """Reconstruct Markdown from segments"""
        lines = []
        for seg in segments:
            if seg.segment_type == 'heading':
                prefix = seg.metadata.get('prefix', '#') if seg.metadata else '#'
                lines.append(f"{prefix} {seg.translated}")
            elif seg.segment_type == 'list_item':
                lines.append(f"- {seg.translated}")
            elif seg.segment_type == 'blockquote':
                lines.append(f"> {seg.translated}")
            else:
                lines.append(seg.translated)
        
        return '\n'.join(lines)


class JSONParser:
    """Parser for JSON files (for translations and config)"""

    @staticmethod
    def parse(file_path: str) -> List[TranslatedSegment]:
        """Parse JSON file extracting translatable strings"""
        segments = []
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        def extract_strings(obj, path=""):
            """Recursively extract strings from JSON"""
            if isinstance(obj, dict):
                for key, value in obj.items():
                    new_path = f"{path}.{key}" if path else key
                    extract_strings(value, new_path)
            elif isinstance(obj, list):
                for idx, item in enumerate(obj):
                    new_path = f"{path}[{idx}]"
                    extract_strings(item, new_path)
            elif isinstance(obj, str) and len(obj) > 2:  # Only translate non-trivial strings
                segment = TranslatedSegment(
                    original=obj,
                    translated="",
                    start_offset=0,
                    end_offset=0,
                    segment_type="json_value",
                    metadata={'path': path}
                )
                segments.append(segment)
        
        extract_strings(data)
        return segments

    @staticmethod
    def reconstruct(segments: List[TranslatedSegment], original_file: str) -> Dict:
        """Reconstruct JSON with translated values"""
        with open(original_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        trans_map = {seg.metadata['path']: seg.translated for seg in segments 
                     if seg.metadata and 'path' in seg.metadata}
        
        def update_values(obj, path=""):
            """Recursively update strings in JSON"""
            if isinstance(obj, dict):
                for key, value in obj.items():
                    new_path = f"{path}.{key}" if path else key
                    if new_path in trans_map:
                        obj[key] = trans_map[new_path]
                    else:
                        update_values(value, new_path)
            elif isinstance(obj, list):
                for idx, item in enumerate(obj):
                    new_path = f"{path}[{idx}]"
                    if new_path in trans_map:
                        obj[idx] = trans_map[new_path]
                    else:
                        update_values(item, new_path)
        
        update_values(data)
        return data


class PDFParser:
    """Parser for PDF documents"""

    @staticmethod
    def parse(file_path: str) -> List[TranslatedSegment]:
        """
        Parse PDF file. 
        Note: Requires pypdf2 or pdfplumber
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber required for PDF support. Install with: pip install pdfplumber")
        
        segments = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    for para in text.split('\n\n'):
                        if para.strip():
                            segment = TranslatedSegment(
                                original=para.strip(),
                                translated="",
                                start_offset=page_num,
                                end_offset=page_num,
                                segment_type="pdf_paragraph",
                                metadata={'page': page_num + 1}
                            )
                            segments.append(segment)
        
        return segments

    @staticmethod
    def reconstruct(segments: List[TranslatedSegment], output_path: str):
        """
        Reconstruct PDF with translations.
        Note: This is a simplified approach - full PDF editing is complex
        """
        # For now, create a text file with translations
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        c = canvas.Canvas(output_path, pagesize=letter)
        y = 750
        
        for seg in segments:
            if seg.translated:
                # Write translated text
                text_obj = c.beginText(50, y)
                text_obj.setFont("Helvetica", 10)
                text_obj.textLines(seg.translated, maxWidth=500)
                c.drawText(text_obj)
                y -= 30
                
                if y < 50:
                    c.showPage()
                    y = 750
        
        c.save()


class DocxParser:
    """Parser for DOCX documents"""

    @staticmethod
    def parse(file_path: str) -> List[TranslatedSegment]:
        """
        Parse DOCX file.
        Note: Requires python-docx
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for DOCX support. Install with: pip install python-docx")
        
        segments = []
        doc = Document(file_path)
        
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                segment = TranslatedSegment(
                    original=para.text,
                    translated="",
                    start_offset=para_idx,
                    end_offset=para_idx,
                    segment_type="docx_paragraph",
                    metadata={'paragraph_index': para_idx, 'style': para.style.name}
                )
                segments.append(segment)
        
        # Also extract tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        segment = TranslatedSegment(
                            original=cell.text,
                            translated="",
                            start_offset=table_idx,
                            end_offset=table_idx,
                            segment_type="docx_table_cell",
                            metadata={'table': table_idx, 'row': row_idx, 'cell': cell_idx}
                        )
                        segments.append(segment)
        
        return segments

    @staticmethod
    def reconstruct(segments: List[TranslatedSegment], original_file: str, output_path: str):
        """Reconstruct DOCX with translations"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required")
        
        doc = Document(original_file)
        
        # Create mapping of original to translated
        trans_map = {}
        for seg in segments:
            trans_map[seg.original] = seg.translated
        
        # Update paragraphs
        for para in doc.paragraphs:
            if para.text in trans_map:
                # Replace entire paragraph
                para.clear()
                para.add_run(trans_map[para.text])
        
        # Update tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text in trans_map:
                        for paragraph in cell.paragraphs:
                            paragraph.clear()
                            paragraph.add_run(trans_map[cell.text])
        
        doc.save(output_path)


class TranslationMemory:
    """Translation memory for consistency across documents"""

    def __init__(self, memory_file: str = "data/translation_memory.json"):
        self.memory_file = memory_file
        self.memory: Dict[str, str] = self._load_memory()

    def _load_memory(self) -> Dict[str, str]:
        """Load translation memory from file"""
        if os.path.exists(self.memory_file):
            with open(self.memory_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}

    def save(self):
        """Save translation memory to file"""
        os.makedirs(os.path.dirname(self.memory_file), exist_ok=True)
        with open(self.memory_file, 'w', encoding='utf-8') as f:
            json.dump(self.memory, f, ensure_ascii=False, indent=2)

    def add(self, original: str, translated: str):
        """Add translation to memory"""
        if original.strip():
            self.memory[original.strip()] = translated.strip()

    def lookup(self, text: str) -> Optional[str]:
        """Look up translation in memory"""
        text = text.strip()
        if text in self.memory:
            return self.memory[text]
        
        # Fuzzy matching for similar strings
        for key, value in self.memory.items():
            if self._similarity(text, key) > 0.85:
                return value
        
        return None

    @staticmethod
    def _similarity(s1: str, s2: str) -> float:
        """Calculate string similarity"""
        from difflib import SequenceMatcher
        return SequenceMatcher(None, s1, s2).ratio()

    def get_stats(self) -> Dict:
        """Get translation memory statistics"""
        return {
            "total_entries": len(self.memory),
            "avg_entry_size": sum(len(k) + len(v) for k, v in self.memory.items()) / len(self.memory) if self.memory else 0,
        }


if __name__ == "__main__":
    # Example usage
    txt_path = "test.txt"
    # segments = DocumentParser.extract_segments(txt_path)
    # print(f"Extracted {len(segments)} segments")
    print("Document translation module loaded successfully")
